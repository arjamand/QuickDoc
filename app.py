import tkinter as tk
import tkinter
from tkinter import ttk, messagebox, filedialog
import customtkinter
import sounddevice as sd
import numpy as np
import whisper
import os
import time
from datetime import datetime
import win32com.client  # For MS Word integration
import threading
import queue
import warnings
from docx import Document
from PIL import Image  # Import the Image class from PIL
import docx
from docx.shared import Inches
from pdf2image import convert_from_path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import fitz
import time 
import psutil
import tempfile

warnings.filterwarnings("ignore", category=UserWarning)

customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")

# Load Whisper model (outside class for performance)
model = whisper.load_model("base")

# --- Image Settings ---
IMAGE_WIDTH = 400  # Reduced width
IMAGE_HEIGHT = 200 # Reduced height
IMAGE_PATH = 'template.png'  # Update to match actual path

class AudioRecorder:
    def __init__(self, sample_rate=16000):
        self.sample_rate = sample_rate
        self.recording = False
        self.audio_queue = queue.Queue()
        self.thread = None

    def list_microphones(self):
        """List all available audio input devices."""
        devices = sd.query_devices()
        input_devices = [dev for dev in devices if dev['max_input_channels'] > 0]
        return input_devices

    def select_microphone(self, input_devices, selected_mic_index):
        """Select a microphone."""
        return input_devices[selected_mic_index], selected_mic_index

    def audio_callback(self, indata, frames, time, status):
        """Callback to store audio data."""
        if status:
            print(f"Status: {status}")
        if self.recording:
            self.audio_queue.put(indata.copy())

    def start_recording(self, device_index):
        """Start audio recording."""
        self.recording = True
        self.audio_queue = queue.Queue()

        try:
            self.stream = sd.InputStream(
                samplerate=self.sample_rate,
                channels=1,
                device=device_index,
                callback=self.audio_callback
            )
            self.stream.start()
        except sd.PortAudioError as e:
            print(f"Audio input error: {e}")
            self.recording = False

    def stop_recording(self):
        """Stop audio recording and collect audio data."""
        if not self.recording:
            return None

        self.recording = False
        self.stream.stop()

        # Collect audio data from queue
        audio_chunks = []
        while not self.audio_queue.empty():
            audio_chunks.append(self.audio_queue.get())

        if not audio_chunks:
            print("No audio recorded.")
            return None

        # Combine and normalize audio
        audio_data = np.concatenate(audio_chunks, axis=0)
        audio_data = np.squeeze(audio_data).astype(np.float32)
        audio_data = audio_data / np.max(np.abs(audio_data))

        return audio_data

    def transcribe_audio(self, audio_data):
        """Transcribe audio using Whisper."""
        if audio_data is None or len(audio_data) == 0:
            return ""

        try:
            result = model.transcribe(audio_data, language="de", fp16=False)
            return result.get("text", "").strip()
        except Exception as e:
            print(f"Transcription error: {e}")
            return ""


class WhisperASRApp(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # Configure window
        self.title("Whisper ASR Transcription")
        self.geometry("1280x640")  # Increased width to accommodate side-by-side layout
        self.audio_recorder = AudioRecorder()
        
        # Configure grid layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure((0, 1), weight=1)

        # Loading widgets
        self.loading_window = None

        # Create sidebar frame
        self.sidebar_frame = customtkinter.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=4, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(10, weight=1)

        # Consistent button width
        BUTTON_WIDTH = 220
        BUTTON_HEIGHT = 40

        # App title
        self.logo_label = customtkinter.CTkLabel(
            self.sidebar_frame,
            text="Whisper ASR",
            font=customtkinter.CTkFont(size=20, weight="bold")
        )
        self.logo_label.grid(row=0, column=0, padx=15, pady=(20, 10))

        # Microphone selection dropdown
        self.mic_label = customtkinter.CTkLabel(self.sidebar_frame, text="Select Microphone:")
        self.mic_label.grid(row=1, column=0, padx=15, pady=(10, 5))

        # Get available microphones
        self.microphones = self.audio_recorder.list_microphones()
        mic_names = [device['name'] for device in self.microphones]

        # Create a frame to hold the dropdown to prevent resizing
        self.mic_dropdown_frame = customtkinter.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.mic_dropdown_frame.grid(row=2, column=0, padx=15, pady=(0, 10), sticky="ew")

        self.mic_dropdown = customtkinter.CTkOptionMenu(
            self.mic_dropdown_frame,
            values=mic_names,
            width=220,  # Fixed width
            command=self.select_microphone
        )
        self.mic_dropdown.pack(fill="x")

        # Recording controls frame
        self.recording_frame = customtkinter.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.recording_frame.grid(row=3, column=0, padx=15, pady=10, sticky="ew")
        self.recording_frame.grid_columnconfigure(1, weight=1)

        # Recording button
        self.record_button = customtkinter.CTkButton(
            self.recording_frame,
            text="Start Recording",
            command=self.toggle_recording,
            width=BUTTON_WIDTH - 100
        )
        self.record_button.grid(row=0, column=0, padx=(0, 5))

        # Recording time display
        self.record_time_label = customtkinter.CTkLabel(
            self.recording_frame,
            text="00:00",
            font=customtkinter.CTkFont(size=14),
            fg_color=self.record_button.cget("fg_color"),
            corner_radius=self.record_button.cget("corner_radius")
        )
        self.record_time_label.grid(row=0, column=1, sticky="ew")

        # Embed Transcription Button
        self.embed_button = customtkinter.CTkButton(
            self.sidebar_frame,
            text="Embed Transcription",
            command=self.embed_transcription,
            width=BUTTON_WIDTH
        )
        self.embed_button.grid(row=4, column=0, padx=15, pady=10)

        # Print Document Button
        self.print_button = customtkinter.CTkButton(
            self.sidebar_frame,
            text="Print Document",
            command=self.print_document,
            width=BUTTON_WIDTH
        )
        self.print_button.grid(row=5, column=0, padx=15, pady=10)

        # Save Document Button
        self.save_button = customtkinter.CTkButton(
            self.sidebar_frame,
            text="Save Document",
            command=self.save_document,
            width=BUTTON_WIDTH
        )
        self.save_button.grid(row=6, column=0, padx=15, pady=10)

        # Create a frame to hold scaling and appearance options
        self.settings_frame = customtkinter.CTkFrame(self.sidebar_frame)
        self.settings_frame.grid(row=11, column=0, padx=15, pady=10, sticky="ew")
        self.settings_frame.grid_columnconfigure(0, weight=1)

        # Appearance Mode
        self.appearance_mode_label = customtkinter.CTkLabel(self.settings_frame, text="Appearance Mode:", anchor="w")
        self.appearance_mode_label.grid(row=0, column=0, padx=0, pady=(0, 0))
        self.appearance_mode_optionemenu = customtkinter.CTkOptionMenu(
            self.settings_frame,
            values=["Light", "Dark", "System"],
            width=BUTTON_WIDTH,
            command=self.change_appearance_mode_event
        )
        self.appearance_mode_optionemenu.grid(row=1, column=0, padx=0, pady=(0, 10))

        # UI Scaling
        self.scaling_label = customtkinter.CTkLabel(self.settings_frame, text="UI Scaling:", anchor="w")
        self.scaling_label.grid(row=2, column=0, padx=0, pady=(0, 0))
        self.scaling_optionemenu = customtkinter.CTkOptionMenu(
            self.settings_frame,
            values=["80%", "90%", "100%", "110%", "120%"],
            width=BUTTON_WIDTH,
            command=self.change_scaling_event
        )
        self.scaling_optionemenu.grid(row=3, column=0, padx=0, pady=(0, 0))

        # Main area for transcription and document
        self.main_frame = customtkinter.CTkFrame(self)
        self.main_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.main_frame.grid_columnconfigure((0, 1), weight=1)
        self.main_frame.grid_rowconfigure((0, 1), weight=1)

        # Create a frame for transcription
        self.transcription_frame = customtkinter.CTkFrame(self.main_frame)
        self.transcription_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        self.transcription_frame.grid_rowconfigure(1, weight=1)
        self.transcription_frame.grid_columnconfigure(0, weight=1)

        # Transcription area
        self.transcription_label = customtkinter.CTkLabel(self.transcription_frame, text="Transcription")
        self.transcription_label.grid(row=0, column=0, padx=10, pady=(0, 5), sticky="w")

        self.transcription_textbox = customtkinter.CTkTextbox(self.transcription_frame, width=500, height=300)
        self.transcription_textbox.grid(row=1, column=0, padx=10, pady=(0, 10), sticky="nsew")
        self.transcription_textbox.insert("0.0", "Transcription will appear here...")

        # Create a frame for document template
        self.doc_frame = customtkinter.CTkFrame(self.main_frame)
        self.doc_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        self.doc_frame.grid_rowconfigure(2, weight=1)
        self.doc_frame.grid_columnconfigure(0, weight=1)

        # Document template label
        self.doc_label = customtkinter.CTkLabel(self.doc_frame, text="Document Template")
        self.doc_label.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="w")

        # Load and display the image with dynamic resizing
        original_image = Image.open(os.path.join(IMAGE_PATH))
        img_width, img_height = original_image.size
        aspect_ratio = img_width / img_height
        
        # Calculate new dimensions maintaining aspect ratio
        # new_height = 200  # Fixed height
        # new_width = int(new_height * aspect_ratio)

        self.doc_image = customtkinter.CTkImage(
            light_image=original_image, 
            size=(600, 600)
        )
        self.image_label = customtkinter.CTkLabel(self.doc_frame, image=self.doc_image, text="")
        self.image_label.grid(row=1, column=0, padx=10, pady=(0,10), sticky="w")

        # Document textbox
        # self.doc_textbox = customtkinter.CTkTextbox(self.doc_frame, width=500, height=300)
        # self.doc_textbox.grid(row=2, column=0, padx=10, pady=(0, 10), sticky="nsew")

        # self.doc_text = ""  # Store the text of the document template
        # self.load_document_template("template_doc.docx")

        # Set default values
        self.appearance_mode_optionemenu.set("Dark")
        self.scaling_optionemenu.set("100%")

        # Recording state variables
        self.is_recording = False
        self.selected_mic_index = 0
        self.recording_start_time = None
        self.audio_data = []

    def show_loading(self, message="Loading..."):
        if self.loading_window is not None:  # If there's already a loading window, close it before opening a new one
           self.loading_window.destroy()
        self.loading_window = tk.Toplevel(self)
        self.loading_window.title("Please Wait")
        self.loading_window.geometry("300x100")
        self.loading_window.resizable(False, False)
        
        loading_label = tk.Label(self.loading_window, text=message, font=("Arial", 12))
        loading_label.pack(pady=30)
        self.loading_window.transient(self)  # Set the loading window as a transient of the main window
        self.loading_window.grab_set()  # Grab focus to prevent interaction with the main window

        # Make sure the loading window is always on top of the main window
        self.loading_window.attributes("-topmost", True)

    def hide_loading(self):
        if self.loading_window:
            self.loading_window.destroy()
            self.loading_window = None

    def select_microphone(self, selected_mic):
        try:
            self.selected_mic_index = [device['name'] for device in self.microphones].index(selected_mic)
        except ValueError:
            messagebox.showerror("Microphone Error", "Selected microphone is no longer available. Please select a different one.")
            return

    def toggle_recording(self):
        if not self.is_recording:
            self.start_recording()
        else:
            self.stop_recording()

    def start_recording(self):
        self.is_recording = True
        self.record_button.configure(text="Stop Recording")
        self.audio_data = []
        self.recording_start_time = time.time()

        self.audio_recorder.start_recording(self.selected_mic_index)
        self.update_record_time()

    def update_record_time(self):
      if self.is_recording:
          elapsed_time = time.time() - self.recording_start_time
          minutes, seconds = divmod(int(elapsed_time), 60)
          self.record_time_label.configure(text=f"{minutes:02d}:{seconds:02d}")
          self.after(100, self.update_record_time)

    def stop_recording(self):
      self.is_recording = False
      self.record_button.configure(text="Start Recording")
      self.record_time_label.configure(text="00:00")  # Reset time display
      audio_data = self.audio_recorder.stop_recording()
      if audio_data is not None:
        self.transcribe_audio(audio_data)

    def transcribe_audio(self, audio_data):
       self.show_loading(message="Transcribing...")
       threading.Thread(target=self._transcribe_audio_thread, args=(audio_data,), daemon=True).start()
    def _transcribe_audio_thread(self, audio_data):
       transcription = self.audio_recorder.transcribe_audio(audio_data)
       self.after(0, self._update_transcription_ui, transcription)  # Use after to update the UI
    def _update_transcription_ui(self, transcription):
       self.transcription_textbox.delete("0.0", "end")
       self.transcription_textbox.insert("0.0", transcription)
       self.hide_loading()
    def is_file_locked(self, file_path):
        """Check if the file is locked by another process."""
        for proc in psutil.process_iter(attrs=['pid', 'name', 'open_files']):
            for file in proc.info['open_files'] or []:
                if file.path == file_path:
                    return True
        return False
    
    def create_image(self, doc_path, image_filename):
        """Convert a DOCX file to an image without locking issues."""
        try:
            # Step 1: Generate a temporary PDF file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
                temp_pdf_path = temp_pdf.name

            try:
                doc = Document(doc_path)
                c = canvas.Canvas(temp_pdf_path, pagesize=letter)
                width, height = letter

                y_position = height - 50
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    if y_position < 50:  # Create a new page if reaching bottom
                        c.showPage()
                        y_position = height - 50
                    c.drawString(50, y_position, text)
                    y_position -= 15

                c.save()  # Save the PDF
                print(f"PDF saved to: {temp_pdf_path}")

            except Exception as e:
                messagebox.showerror("Error", f"Error creating PDF: {e}")
                return

            # Step 2: Convert PDF to Image using PyMuPDF (fitz)
            try:
                pdf_document = fitz.open(temp_pdf_path)  # Open the temporary PDF
                first_page = pdf_document.load_page(0)   # Load the first page
                pix = first_page.get_pixmap()
                pix.save(image_filename)  # Save the image
                print(f"Image saved successfully at: {image_filename}")
                pdf_document.close()

            except Exception as e:
                messagebox.showerror("Error", f"Error converting PDF to image: {e}")
                return

            # Step 3: Clean up temporary PDF file
            if os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
                print("Temporary PDF file deleted.")

        except Exception as e:
            messagebox.showerror("Error", f"Error during conversion: {e}")
        # """Helper method to convert the new document to an image."""
        
        # try:
        #     #first convert to pdf
        #     print(doc_path)
        #     pdf_path = os.path.splitext(doc_path)[0] + ".pdf"
        #     print(pdf_path)
            
        #     # Create a temporary PDF from the docx to convert into images
        #     import subprocess
        #     try:
        #       subprocess.run(["soffice", "--headless", "--convert-to", "pdf", doc_path, "--outdir", os.path.dirname(pdf_path)], check=True)
        #     except Exception as e:
        #        messagebox.showerror("Error", f"Error converting to PDF: {e}")
        #        return
            
            
        #     images = convert_from_path(pdf_path)  # returns a list of PIL Image objects
            
        #     # For a simple case we will just take the first page
        #     if images:
        #         first_page_image = images[0]
        #         first_page_image_path = image_filename
        #         first_page_image.save(first_page_image_path)
        #         print(f"Image saved: {first_page_image_path}")
            
            
        #     #cleanup temp pdf
        #     os.remove(pdf_path)
            
        # except Exception as e:
        #     messagebox.showerror("Error", f"Error converting to Image: {e}")

    
    def embed_transcription(self):
        transcription = self.transcription_textbox.get("0.0", "end").strip()
        if not transcription:
            messagebox.showwarning("Warning", "Please enter transcription text.")
            return

        try:
            # 1. Open Template
            template_path = "template.docx"  # Replace with actual path if needed
            doc = docx.Document(template_path)

            # 2. Add Text to Body (at the end)
            doc.add_paragraph(transcription) # Add a paragraph with the transcription
            
            # 3. Save New Doc
            prescriptions_folder = "prescriptions"
            if not os.path.exists(prescriptions_folder):
                os.makedirs(prescriptions_folder)

            new_filename = "updated_prescription.docx" #You might want a more unique name
            new_doc_path = os.path.join(prescriptions_folder, new_filename)

            doc.save(new_doc_path)
            
            # 4. Create Image (will save in root folder with a similar name)
            self.create_image(new_doc_path, os.path.splitext(new_filename)[0] + ".png")

            messagebox.showinfo("Success", "Document updated and saved.")


        except FileNotFoundError:
            messagebox.showerror("Error", f"Template file not found: {template_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    # def embed_transcription(self):
    #     transcription = self.transcription_textbox.get("0.0", "end")
    #     updated_doc = self.doc_text + "\n\n" + transcription
    #     self.doc_textbox.delete("0.0", "end")
    #     self.doc_textbox.insert("0.0", updated_doc)

    def load_document_template(self, doc_path):
        try:
            doc = Document(doc_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            self.doc_text = text  # Store the document text
            self.doc_textbox.delete("0.0", "end")
            self.doc_textbox.insert("0.0", self.doc_text)
        except Exception as e:
             self.doc_textbox.delete("0.0", "end")
             self.doc_textbox.insert("0.0", f"Error loading document: {e}")

    def print_document(self):
          try:
             # Open MS Word
             word = win32com.client.Dispatch("Word.Application")
             word.Visible = False
             doc = word.Documents.Add()

             # Get document content
             content = self.doc_textbox.get("0.0", "end")
             doc.Content.Text = content
             doc.PrintOut()
             # Close document and word application
             doc.Close(False)
             word.Quit()

          except Exception as e:
            tkinter.messagebox.showerror("Printing Error", str(e))
    def save_document(self):
        try:
            # Open MS Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Add()

            # Get document content
            content = self.doc_textbox.get("0.0", "end")
            doc.Content.Text = content

            # Save document
            filename = f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.SaveAs(filename)

            # Close document and word application
            doc.Close()
            word.Quit()

            tkinter.messagebox.showinfo("Save Successful", f"Document saved as {filename}")
        except Exception as e:
            tkinter.messagebox.showerror("Save Error", str(e))

    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def change_scaling_event(self, new_scaling: str):
        new_scaling_float = int(new_scaling.replace("%", "")) / 100
        customtkinter.set_widget_scaling(new_scaling_float)

def main():
    app = WhisperASRApp()
    app.mainloop()


if __name__ == "__main__":
    main()

